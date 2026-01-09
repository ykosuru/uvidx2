"""
Document Parser - Extracts indexable chunks from PDFs and documents

Supports PDF, DOCX, and plain text documents with layout-aware
chunking that preserves document structure.
"""

import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass

from .base import ContentParser
from ..models import (
    IndexableChunk,
    SourceType,
    SemanticType,
    SourceReference,
    DomainMatch
)
from ..vocabulary import DomainVocabulary


@dataclass
class DocumentSection:
    """Represents a section of a document"""
    title: Optional[str]
    content: str
    page_number: Optional[int]
    section_level: int  # 0 = top level, 1 = subsection, etc.
    element_type: str  # 'heading', 'paragraph', 'table', 'list'


class DocumentParser(ContentParser):
    """
    Parser for document files (PDF, DOCX, TXT, MD)
    
    Uses layout-aware extraction to preserve document structure
    and create meaningful chunks based on sections and headings.
    """
    
    SOURCE_TYPE = SourceType.DOCUMENT
    SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.doc', '.txt', '.md', '.html']
    
    def __init__(self, vocabulary: DomainVocabulary):
        """Initialize document parser"""
        super().__init__(vocabulary)
        
        # Try to import optional PDF libraries
        self._pdf_library = None
        self._try_load_pdf_libraries()
    
    def _try_load_pdf_libraries(self):
        """Try to load PDF parsing libraries"""
        # Try pdfplumber first (best for tables)
        try:
            import pdfplumber
            self._pdf_library = 'pdfplumber'
            return
        except ImportError:
            pass
        
        # Try PyMuPDF (fast)
        try:
            import fitz  # PyMuPDF
            self._pdf_library = 'pymupdf'
            return
        except ImportError:
            pass
        
        # Try pypdf (basic)
        try:
            from pypdf import PdfReader
            self._pdf_library = 'pypdf'
            return
        except ImportError:
            pass
        
        print("Warning: No PDF library found. Install pdfplumber, pymupdf, or pypdf")
    
    def can_parse(self, file_path: str) -> bool:
        """Check if file is a supported document"""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        # PDF requires library
        if ext == '.pdf' and not self._pdf_library:
            return False
        
        return ext in self.SUPPORTED_EXTENSIONS
    
    def parse(self, content: bytes, source_path: str) -> List[IndexableChunk]:
        """Parse document content and extract chunks"""
        path = Path(source_path)
        ext = path.suffix.lower()
        
        if ext == '.pdf':
            return self._parse_pdf(content, source_path)
        elif ext in ['.docx', '.doc']:
            return self._parse_docx(content, source_path)
        elif ext == '.md':
            return self._parse_markdown(content, source_path)
        elif ext == '.html':
            return self._parse_html(content, source_path)
        else:
            return self._parse_text(content, source_path)
    
    def _parse_pdf(self, content: bytes, source_path: str) -> List[IndexableChunk]:
        """Parse PDF document"""
        if self._pdf_library == 'pdfplumber':
            return self._parse_pdf_pdfplumber(content, source_path)
        elif self._pdf_library == 'pymupdf':
            return self._parse_pdf_pymupdf(content, source_path)
        elif self._pdf_library == 'pypdf':
            return self._parse_pdf_pypdf(content, source_path)
        else:
            raise RuntimeError("No PDF library available")
    
    def _parse_pdf_pdfplumber(self, content: bytes, source_path: str) -> List[IndexableChunk]:
        """Parse PDF using pdfplumber"""
        import pdfplumber
        import io
        
        chunks = []
        sections = []
        
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text
                text = page.extract_text() or ""
                
                if text.strip():
                    # Try to identify sections by font size or formatting
                    lines = text.split('\n')
                    current_section = None
                    current_content = []
                    
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # Heuristic: lines that look like headings
                        if self._is_heading(line):
                            # Save previous section
                            if current_content:
                                sections.append(DocumentSection(
                                    title=current_section,
                                    content='\n'.join(current_content),
                                    page_number=page_num,
                                    section_level=0,
                                    element_type='paragraph'
                                ))
                            current_section = line
                            current_content = []
                        else:
                            current_content.append(line)
                    
                    # Add remaining content
                    if current_content:
                        sections.append(DocumentSection(
                            title=current_section,
                            content='\n'.join(current_content),
                            page_number=page_num,
                            section_level=0,
                            element_type='paragraph'
                        ))
                
                # Extract tables
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        table_text = self._format_table(table)
                        sections.append(DocumentSection(
                            title=None,
                            content=table_text,
                            page_number=page_num,
                            section_level=0,
                            element_type='table'
                        ))
        
        # Convert sections to chunks
        chunks = self._sections_to_chunks(sections, source_path)
        return chunks
    
    def _parse_pdf_pymupdf(self, content: bytes, source_path: str) -> List[IndexableChunk]:
        """Parse PDF using PyMuPDF"""
        import fitz
        
        sections = []
        doc = fitz.open(stream=content, filetype="pdf")
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            
            if text.strip():
                # Get text blocks with positioning
                blocks = page.get_text("blocks")
                
                for block in blocks:
                    if block[6] == 0:  # Text block
                        block_text = block[4].strip()
                        if block_text:
                            element_type = 'heading' if self._is_heading(block_text) else 'paragraph'
                            sections.append(DocumentSection(
                                title=block_text if element_type == 'heading' else None,
                                content=block_text,
                                page_number=page_num + 1,
                                section_level=0,
                                element_type=element_type
                            ))
        
        doc.close()
        return self._sections_to_chunks(sections, source_path)
    
    def _parse_pdf_pypdf(self, content: bytes, source_path: str) -> List[IndexableChunk]:
        """Parse PDF using pypdf"""
        from pypdf import PdfReader
        import io
        
        sections = []
        reader = PdfReader(io.BytesIO(content))
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            
            if text.strip():
                sections.append(DocumentSection(
                    title=None,
                    content=text,
                    page_number=page_num,
                    section_level=0,
                    element_type='paragraph'
                ))
        
        return self._sections_to_chunks(sections, source_path)
    
    def _parse_docx(self, content: bytes, source_path: str) -> List[IndexableChunk]:
        """Parse DOCX document"""
        try:
            from docx import Document
            import io
        except ImportError:
            print("Warning: python-docx not installed, treating as text")
            return self._parse_text(content, source_path)
        
        sections = []
        doc = Document(io.BytesIO(content))
        
        current_section_title = None
        current_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a heading
            if para.style.name.startswith('Heading'):
                # Save previous section
                if current_content:
                    sections.append(DocumentSection(
                        title=current_section_title,
                        content='\n'.join(current_content),
                        page_number=None,
                        section_level=0,
                        element_type='paragraph'
                    ))
                current_section_title = text
                current_content = []
            else:
                current_content.append(text)
        
        # Add remaining content
        if current_content:
            sections.append(DocumentSection(
                title=current_section_title,
                content='\n'.join(current_content),
                page_number=None,
                section_level=0,
                element_type='paragraph'
            ))
        
        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            
            if rows:
                table_text = self._format_table(rows)
                sections.append(DocumentSection(
                    title=None,
                    content=table_text,
                    page_number=None,
                    section_level=0,
                    element_type='table'
                ))
        
        return self._sections_to_chunks(sections, source_path)
    
    def _parse_markdown(self, content: bytes, source_path: str) -> List[IndexableChunk]:
        """Parse Markdown document"""
        text = content.decode('utf-8', errors='replace')
        sections = []
        
        # Split by headers
        header_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
        
        # Find all headers
        headers = [(m.start(), len(m.group(1)), m.group(2)) for m in header_pattern.finditer(text)]
        
        if not headers:
            # No headers, treat as single section
            sections.append(DocumentSection(
                title=None,
                content=text,
                page_number=None,
                section_level=0,
                element_type='paragraph'
            ))
        else:
            # Split by headers
            for i, (pos, level, title) in enumerate(headers):
                # Get content until next header
                if i + 1 < len(headers):
                    end_pos = headers[i + 1][0]
                else:
                    end_pos = len(text)
                
                content_text = text[pos:end_pos]
                # Remove the header line from content
                content_text = header_pattern.sub('', content_text, count=1).strip()
                
                if content_text:
                    sections.append(DocumentSection(
                        title=title,
                        content=content_text,
                        page_number=None,
                        section_level=level - 1,
                        element_type='section'
                    ))
        
        return self._sections_to_chunks(sections, source_path)
    
    def _parse_html(self, content: bytes, source_path: str) -> List[IndexableChunk]:
        """Parse HTML document"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            # Fallback: strip tags with regex
            text = content.decode('utf-8', errors='replace')
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text)
            return self._parse_text(text.encode('utf-8'), source_path)
        
        soup = BeautifulSoup(content, 'html.parser')
        sections = []
        
        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()
        
        # Find headings and their content
        for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            level = int(heading.name[1])
            title = heading.get_text(strip=True)
            
            # Get content until next heading
            content_parts = []
            for sibling in heading.find_next_siblings():
                if sibling.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    break
                content_parts.append(sibling.get_text(strip=True))
            
            if content_parts:
                sections.append(DocumentSection(
                    title=title,
                    content='\n'.join(content_parts),
                    page_number=None,
                    section_level=level - 1,
                    element_type='section'
                ))
        
        # If no sections found, get all text
        if not sections:
            text = soup.get_text(separator='\n', strip=True)
            sections.append(DocumentSection(
                title=None,
                content=text,
                page_number=None,
                section_level=0,
                element_type='paragraph'
            ))
        
        return self._sections_to_chunks(sections, source_path)
    
    def _parse_text(self, content: bytes, source_path: str) -> List[IndexableChunk]:
        """Parse plain text document"""
        text = content.decode('utf-8', errors='replace')
        
        # Try to find natural section breaks
        sections = []
        
        # Split by double newlines (paragraph breaks)
        paragraphs = re.split(r'\n\n+', text)
        
        current_section_title = None
        current_content = []
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Check if paragraph looks like a heading
            lines = para.split('\n')
            if len(lines) == 1 and self._is_heading(para):
                # Save previous section
                if current_content:
                    sections.append(DocumentSection(
                        title=current_section_title,
                        content='\n\n'.join(current_content),
                        page_number=None,
                        section_level=0,
                        element_type='paragraph'
                    ))
                current_section_title = para
                current_content = []
            else:
                current_content.append(para)
        
        # Add remaining content
        if current_content:
            sections.append(DocumentSection(
                title=current_section_title,
                content='\n\n'.join(current_content),
                page_number=None,
                section_level=0,
                element_type='paragraph'
            ))
        
        return self._sections_to_chunks(sections, source_path)
    
    def _is_heading(self, text: str) -> bool:
        """Heuristic to detect if text looks like a heading"""
        text = text.strip()
        
        if not text:
            return False
        
        # Short lines that end without punctuation
        if len(text) < 100 and not text.endswith(('.', ',', ';', ':')):
            # Check for heading patterns
            if text.isupper():
                return True
            if re.match(r'^\d+\.?\s+\w', text):  # Numbered heading
                return True
            if re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*$', text):  # Title Case
                return True
            # All caps with possible numbers
            if re.match(r'^[A-Z0-9\s\-:]+$', text) and len(text) < 80:
                return True
        
        return False
    
    def _format_table(self, table: List[List[str]]) -> str:
        """Format a table as text"""
        if not table:
            return ""
        
        lines = []
        for row in table:
            if row:
                lines.append(' | '.join(str(cell or '') for cell in row))
        
        return '\n'.join(lines)
    
    def _sections_to_chunks(self, 
                            sections: List[DocumentSection],
                            source_path: str) -> List[IndexableChunk]:
        """Convert document sections to indexable chunks"""
        chunks = []
        
        for i, section in enumerate(sections):
            content = section.content.strip()
            if not content:
                continue
            
            # Skip very short content
            if len(content) < 20:
                continue
            
            # Chunk long sections
            if len(content) > 1500:
                text_chunks = self.chunk_text(content, max_chunk_size=1200, overlap=100)
                for j, chunk_text in enumerate(text_chunks):
                    chunk = self._create_section_chunk(
                        section=section,
                        content=chunk_text,
                        source_path=source_path,
                        chunk_index=f"{i}_{j}"
                    )
                    chunks.append(chunk)
            else:
                chunk = self._create_section_chunk(
                    section=section,
                    content=content,
                    source_path=source_path,
                    chunk_index=str(i)
                )
                chunks.append(chunk)
        
        return chunks
    
    def _create_section_chunk(self,
                               section: DocumentSection,
                               content: str,
                               source_path: str,
                               chunk_index: str) -> IndexableChunk:
        """Create a chunk from a document section"""
        
        # Match domain concepts
        domain_matches = self.match_domain_concepts(content)
        
        # Determine semantic type
        if section.element_type == 'table':
            semantic_type = SemanticType.TABLE
        elif section.element_type == 'heading':
            semantic_type = SemanticType.HEADING
        elif section.element_type == 'list':
            semantic_type = SemanticType.LIST
        elif section.title:
            semantic_type = SemanticType.SECTION
        else:
            semantic_type = SemanticType.PARAGRAPH
        
        # Build metadata
        metadata = {
            'element_type': section.element_type,
            'section_level': section.section_level
        }
        if section.title:
            metadata['section_title'] = section.title
        
        # Create source reference
        source_ref = SourceReference(
            file_path=source_path,
            page_number=section.page_number,
            section_title=section.title
        )
        
        # Create embedding text
        embedding_text = self.create_embedding_text(
            content,
            semantic_type,
            domain_matches,
            metadata
        )
        
        return IndexableChunk(
            chunk_id=self.generate_chunk_id(source_path, content, chunk_index),
            text=content,
            embedding_text=embedding_text,
            source_type=SourceType.DOCUMENT,
            semantic_type=semantic_type,
            source_ref=source_ref,
            domain_matches=domain_matches,
            metadata=metadata
        )
